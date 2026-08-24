"""Editable Word projection for website copy and its canonical JSON sources.

The website uses small manifest-backed JSON files plus ``$source`` references
to canonical facts and image descriptions.  This module exposes those authored
strings as stable Word Content Controls, shows the matching website typography
beside every value, and writes edits back to the owning JSON file without
flattening any reference objects.
"""

from __future__ import annotations

from collections import OrderedDict
from dataclasses import dataclass, field
import hashlib
import re
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from content_model import read_json_value, write_json_atomic
from design_tokens import DesignTokens, TextStyle, load_design_tokens
from project_paths import (
    ASSET_RECORD_PATH,
    CONTENT_DIR,
    DESIGN_TOKENS_PATH,
    FACTS_CONTENT_PATH,
    ROOT,
    SITE_CONTENT_PATH,
    WEBSITE_WORKING_DOCX_PATH,
)
from resume.word_renderer import ContentControlNotFoundError, read_content_control_values


PathPart = str | int
JsonPath = tuple[PathPart, ...]
WEBSITE_TAG_PREFIX = "WEB_"
DOCUMENT_VERSION = 1
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
TABLE_WIDTHS_DXA = (2700, 6660)
CELL_MARGIN_VERTICAL_DXA = 80
CELL_MARGIN_HORIZONTAL_DXA = 120
PREVIEW_SIZE_CAP_PT = 34.0


class WebsiteWorkingError(RuntimeError):
    """Raised when the website working document cannot round-trip safely."""


@dataclass
class WebsiteContentField:
    """One editable canonical string and its first website presentation."""

    tag: str
    label: str
    source_path: Path
    json_path: JsonPath
    value: str
    group: str
    style_name: str
    context: str
    template: str | None = None
    usages: list[str] = field(default_factory=list)

    @property
    def source_label(self) -> str:
        return f"{self.source_path.relative_to(ROOT).as_posix()} : {_format_json_path(self.json_path)}"


@dataclass(frozen=True)
class WebsiteSourceUpdate:
    """One fully updated JSON source ready for an atomic save."""

    path: Path
    value: Any


@dataclass(frozen=True)
class _SourceRecord:
    path: Path
    value: Any
    json_path: JsonPath
    context_path: JsonPath
    group: str
    section_id: str


_SKIPPED_BRANCH_KEYS = {
    "$source",
    "display",
    "external",
    "file",
    "href",
    "id",
    "include",
    "resume_id",
    "resume_ids",
    "resume_path",
    "social_image",
    "src",
    "url",
    "wide",
}


def collect_website_fields() -> list[WebsiteContentField]:
    """Return every editable website string in page order.

    Direct strings remain owned by their section file.  A ``$source`` object is
    followed to ``facts.json``, ``asset-record.json``, or ``site.json`` and the
    canonical leaf is exposed only once even when the page reuses it.
    """

    sources = _load_source_documents()
    collector = _FieldCollector(sources)
    for record in _website_source_records(sources):
        collector.visit(record)
    return list(collector.fields.values())


def create_working_website(output_path: Path = WEBSITE_WORKING_DOCX_PATH) -> None:
    """Create a populated Word editor for all authored website text."""

    fields = collect_website_fields()
    if not fields:
        raise WebsiteWorkingError("No editable website fields were discovered")
    tokens = load_design_tokens(DESIGN_TOKENS_PATH)
    document = Document()
    _configure_document(document, tokens)
    _add_document_intro(document, len(fields), tokens)

    grouped: OrderedDict[str, list[WebsiteContentField]] = OrderedDict()
    for content_field in fields:
        grouped.setdefault(content_field.group, []).append(content_field)
    for group, group_fields in grouped.items():
        _add_group(document, group, group_fields, tokens)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    validate_website_working_document(output_path)


def validate_website_working_document(path: Path = WEBSITE_WORKING_DOCX_PATH) -> None:
    """Require exactly one current Content Control for every canonical field."""

    expected = {content_field.tag for content_field in collect_website_fields()}
    try:
        values = read_content_control_values(path)
    except (FileNotFoundError, ContentControlNotFoundError) as error:
        raise WebsiteWorkingError(str(error)) from error
    actual = {tag for tag in values if tag.startswith(WEBSITE_TAG_PREFIX)}
    missing = sorted(expected - actual)
    stale = sorted(actual - expected)
    if missing or stale:
        details: list[str] = []
        if missing:
            details.append(f"missing {len(missing)} current field(s)")
        if stale:
            details.append(f"contains {len(stale)} stale field(s)")
        raise WebsiteWorkingError(
            "website-working.docx " + " and ".join(details) + "; sync or recreate it before changing the JSON schema"
        )


def project_website_updates(values: dict[str, str]) -> list[WebsiteSourceUpdate]:
    """Project Word values into in-memory JSON sources without writing them."""

    fields = collect_website_fields()
    expected = {content_field.tag for content_field in fields}
    supplied = {tag for tag in values if tag.startswith(WEBSITE_TAG_PREFIX)}
    missing = sorted(expected - supplied)
    stale = sorted(supplied - expected)
    if missing or stale:
        parts = []
        if missing:
            parts.append("missing " + ", ".join(missing))
        if stale:
            parts.append("stale " + ", ".join(stale))
        raise WebsiteWorkingError("Website Content Control inventory mismatch: " + "; ".join(parts))

    documents = _load_source_documents()
    changed_paths: set[Path] = set()
    for content_field in fields:
        value = values[content_field.tag].strip()
        _set_json_path(documents[content_field.source_path], content_field.json_path, value)
        changed_paths.add(content_field.source_path)
    return [WebsiteSourceUpdate(path, documents[path]) for path in documents if path in changed_paths]


def write_website_updates(updates: Iterable[WebsiteSourceUpdate]) -> None:
    """Persist a validated website projection with per-file atomic replaces."""

    for update in updates:
        write_json_atomic(update.path, update.value)


def read_working_website_updates(path: Path = WEBSITE_WORKING_DOCX_PATH) -> list[WebsiteSourceUpdate]:
    """Read and validate the working document, returning staged JSON updates."""

    validate_website_working_document(path)
    return project_website_updates(read_content_control_values(path))


class _FieldCollector:
    def __init__(self, documents: OrderedDict[Path, Any]) -> None:
        self.documents = documents
        self.fields: OrderedDict[tuple[Path, JsonPath], WebsiteContentField] = OrderedDict()
        self.tags: set[str] = set()

    def visit(self, record: _SourceRecord) -> None:
        self._visit_value(
            record.value,
            record.path,
            record.json_path,
            record.context_path,
            record.group,
            record.section_id,
            template=None,
        )

    def _visit_value(
        self,
        value: Any,
        source_path: Path,
        json_path: JsonPath,
        context_path: JsonPath,
        group: str,
        section_id: str,
        template: str | None,
    ) -> None:
        if isinstance(value, dict) and "$source" in value:
            source = value.get("$source")
            if not isinstance(source, str):
                raise WebsiteWorkingError(f"Invalid $source at {_format_json_path(context_path)}")
            resolved_path, resolved_json_path = _resolve_source_locator(source)
            resolved_value = _get_json_path(self.documents[resolved_path], resolved_json_path)
            source_template = value.get("$template")
            if source_template is not None and not isinstance(source_template, str):
                raise WebsiteWorkingError(f"Invalid $template for {source}")
            self._visit_value(
                resolved_value,
                resolved_path,
                resolved_json_path,
                context_path,
                group,
                section_id,
                source_template,
            )
            return
        if isinstance(value, dict):
            for key, child in value.items():
                if key.startswith("$") or key in _SKIPPED_BRANCH_KEYS:
                    continue
                self._visit_value(
                    child,
                    source_path,
                    (*json_path, key),
                    (*context_path, key),
                    group,
                    section_id,
                    template,
                )
            return
        if isinstance(value, list):
            for index, child in enumerate(value):
                self._visit_value(
                    child,
                    source_path,
                    (*json_path, index),
                    (*context_path, index),
                    group,
                    section_id,
                    template,
                )
            return
        if not isinstance(value, str):
            return

        locator = (source_path, json_path)
        usage = _format_json_path(context_path)
        existing = self.fields.get(locator)
        if existing is not None:
            if usage not in existing.usages:
                existing.usages.append(usage)
            return
        tag = _field_tag(source_path, json_path)
        if tag in self.tags:
            raise WebsiteWorkingError(f"Website field tag collision: {tag}")
        self.tags.add(tag)
        self.fields[locator] = WebsiteContentField(
            tag=tag,
            label=_field_label(context_path),
            source_path=source_path,
            json_path=json_path,
            value=value,
            group=group,
            style_name=_infer_style(section_id, context_path),
            context=usage,
            template=template,
            usages=[usage],
        )


def _load_source_documents() -> OrderedDict[Path, Any]:
    paths: list[Path] = [SITE_CONTENT_PATH, FACTS_CONTENT_PATH, ASSET_RECORD_PATH]
    manifest = read_json_value(SITE_CONTENT_PATH)
    sections = manifest.get("website", {}).get("sections", []) if isinstance(manifest, dict) else []
    for section in sections:
        if not isinstance(section, dict):
            continue
        if isinstance(section.get("file"), str):
            paths.append(_content_path(section["file"]))
        for item in section.get("items", []):
            if isinstance(item, dict) and isinstance(item.get("file"), str):
                paths.append(_content_path(item["file"]))
    documents: OrderedDict[Path, Any] = OrderedDict()
    for path in paths:
        if path not in documents:
            documents[path] = read_json_value(path)
    return documents


def _website_source_records(documents: OrderedDict[Path, Any]) -> list[_SourceRecord]:
    records: list[_SourceRecord] = []
    site = documents[SITE_CONTENT_PATH]
    global_site = {key: value for key, value in site.items() if key != "website"}
    records.append(
        _SourceRecord(SITE_CONTENT_PATH, global_site, (), ("site_source",), "Site settings & navigation", "site")
    )

    manifest = documents[SITE_CONTENT_PATH]
    sections = manifest.get("website", {}).get("sections", []) if isinstance(manifest, dict) else []
    for section_index, section in enumerate(sections):
        if not isinstance(section, dict) or not isinstance(section.get("id"), str):
            continue
        section_id = section["id"]
        group = _group_label(section_id)
        metadata = {key: value for key, value in section.items() if key not in {"id", "file", "items"}}
        if metadata:
            for key, value in metadata.items():
                records.append(
                    _SourceRecord(
                        SITE_CONTENT_PATH,
                        value,
                        ("website", "sections", section_index, key),
                        ("website", section_id, key),
                        group,
                        section_id,
                    )
                )
        if isinstance(section.get("file"), str):
            path = _content_path(section["file"])
            records.append(_SourceRecord(path, documents[path], (), ("website", section_id), group, section_id))
        for item in section.get("items", []):
            if not isinstance(item, dict) or not isinstance(item.get("file"), str):
                continue
            path = _content_path(item["file"])
            item_id = str(item.get("id", path.stem))
            records.append(
                _SourceRecord(
                    path,
                    documents[path],
                    (),
                    ("website", section_id, item_id),
                    f"{group} / {item_id.replace('_', ' ').title()}",
                    section_id,
                )
            )
    return records


def _content_path(relative_path: str) -> Path:
    path = (CONTENT_DIR / relative_path).resolve()
    try:
        path.relative_to(CONTENT_DIR.resolve())
    except ValueError as error:
        raise WebsiteWorkingError(f"Website source is outside content/: {relative_path}") from error
    return path


def _resolve_source_locator(source: str) -> tuple[Path, JsonPath]:
    if source.startswith("facts."):
        return FACTS_CONTENT_PATH, tuple(source.removeprefix("facts.").split("."))
    if source.startswith("assets."):
        return ASSET_RECORD_PATH, tuple(source.removeprefix("assets.").split("."))
    return SITE_CONTENT_PATH, tuple(source.split("."))


def _get_json_path(value: Any, json_path: JsonPath) -> Any:
    current = value
    for part in json_path:
        try:
            current = current[part]
        except (KeyError, IndexError, TypeError) as error:
            raise WebsiteWorkingError(f"Missing website source: {_format_json_path(json_path)}") from error
    return current


def _set_json_path(value: Any, json_path: JsonPath, replacement: str) -> None:
    if not json_path:
        raise WebsiteWorkingError("A whole JSON source cannot be replaced by a text control")
    parent = _get_json_path(value, json_path[:-1])
    part = json_path[-1]
    try:
        current = parent[part]
    except (KeyError, IndexError, TypeError) as error:
        raise WebsiteWorkingError(f"Missing website source: {_format_json_path(json_path)}") from error
    if not isinstance(current, str):
        raise WebsiteWorkingError(f"Website source is no longer text: {_format_json_path(json_path)}")
    parent[part] = replacement


def _field_tag(source_path: Path, json_path: JsonPath) -> str:
    relative_source = source_path.relative_to(ROOT).as_posix()
    # The manifest moved from details.json into site.json. Keep the established
    # control IDs for its metadata fields so existing Word edits still import.
    if source_path == SITE_CONTENT_PATH and json_path[:2] == ("website", "sections"):
        relative_source = "content/details.json"
    # Asset metadata moved below content/. Keep its established control IDs so
    # existing Word edits continue to target the same canonical fields.
    if source_path == ASSET_RECORD_PATH:
        relative_source = "assets/asset-record.json"
    locator = f"{relative_source}:{_format_json_path(json_path)}"
    digest = hashlib.sha256(locator.encode("utf-8")).hexdigest()[:20].upper()
    return WEBSITE_TAG_PREFIX + digest


def _format_json_path(json_path: JsonPath) -> str:
    result = ""
    for part in json_path:
        if isinstance(part, int):
            result += f"[{part}]"
        else:
            result += ("." if result else "") + part
    return result or "$"


def _field_label(context_path: JsonPath) -> str:
    strings = [part for part in context_path if isinstance(part, str)]
    label = strings[-1] if strings else "Text"
    index = next((part for part in reversed(context_path) if isinstance(part, int)), None)
    friendly = label.replace("_", " ").strip().title()
    return f"{friendly} {index + 1}" if index is not None and label not in {"site_source"} else friendly


def _group_label(section_id: str) -> str:
    return {
        "hero": "00 / Hero",
        "profile": "01 / Engineering profile",
        "case_studies": "02-04 / Featured engineering projects",
        "experience": "05 / Engineering experience",
        "skills": "06 / Technical skills",
        "documentation": "07 / Engineering evidence",
        "leadership": "08 / Leadership & communication",
        "personal_builds": "09 / Hands-on learning",
        "contact": "10 / Contact",
    }.get(section_id, section_id.replace("_", " ").title())


def _infer_style(section_id: str, context_path: JsonPath) -> str:
    parts = [str(part) for part in context_path]
    last = next((part for part in reversed(parts) if not part.isdigit()), "")
    joined = ".".join(parts)

    if section_id == "site":
        if "navigation" in parts and last == "label":
            return "site.navigation"
        if last == "initials":
            return "site.brand"
        if last in {"meta_description", "social_description", "social_image_alt"}:
            return "site.body_muted"
        if last in {"title", "social_title"}:
            return "site.title_small"
        if last in {"email", "website", "phone", "location", "name"}:
            return "site.contact_link"
        return "site.body"
    if section_id == "hero":
        if last in {"eyebrow", "label"}:
            return "site.label"
        if last == "title" and "proof_points" not in joined and "image" not in joined:
            return "site.hero_title"
        if last == "description":
            return "site.hero_intro"
        if "facts" in parts or last == "detail":
            return "site.fact"
        if last == "value":
            return "site.title_small"
        return "site.figure_note" if last in {"alt", "title", "caption"} else "site.body"
    if section_id == "profile":
        return {
            "number": "site.label",
            "heading": "site.title",
            "lead": "site.lead",
            "body": "site.body_muted",
            "label": "site.metadata_label",
            "text": "site.body_detail",
        }.get(last, "site.body")
    if section_id == "case_studies":
        if last in {"number", "type", "label"}:
            return "site.label" if "metadata" not in parts else "site.metadata_label"
        if last == "title":
            if "decision" in parts:
                return "site.callout_title"
            if "image" in parts or "figures" in parts:
                return "site.figure_caption"
            return "site.case_title"
        if last == "subtitle":
            return "site.subtitle"
        if last == "value":
            return "site.metadata_value"
        if last == "process":
            return "site.process"
        if last == "tools":
            return "site.tag"
        if last in {"alt", "caption"}:
            return "site.figure_note"
        return "site.body_detail"
    if section_id == "experience":
        return {
            "number": "site.label",
            "heading": "site.title",
            "date": "site.experience_date",
            "role": "site.experience_title",
            "organization": "site.experience_meta",
            "text": "site.experience_body",
            "tags": "site.tag",
        }.get(last, "site.experience_body")
    if section_id == "skills":
        return {
            "number": "site.label",
            "heading": "site.title",
            "title": "site.skill_title",
            "items": "site.skill_item",
        }.get(last, "site.skill_item")
    if section_id == "documentation":
        return {
            "number": "site.label",
            "heading": "site.title",
            "intro": "site.subtitle",
            "figure": "site.card_label",
            "title": "site.card_title",
            "caption": "site.card_body",
            "alt": "site.figure_note",
        }.get(last, "site.card_body")
    if section_id == "leadership":
        return {
            "number": "site.label",
            "heading": "site.title",
            "title": "site.leadership_title",
            "text": "site.experience_body",
        }.get(last, "site.experience_body")
    if section_id == "personal_builds":
        return {
            "number": "site.label",
            "heading": "site.title",
            "title": "site.personal_title",
            "text": "site.skill_item",
            "alt": "site.figure_note",
        }.get(last, "site.skill_item")
    if section_id == "contact":
        return {
            "number": "site.label",
            "heading": "site.contact_title",
            "lead": "site.contact_lead",
            "label": "site.contact_link",
        }.get(last, "site.contact_link")
    return "site.body"


def _configure_document(document: Document, tokens: DesignTokens) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    _style_font(normal, "Calibri", 11, "1F2933")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = (
        ("Heading 1", 16, 18, 10),
        ("Heading 2", 13, 14, 7),
        ("Heading 3", 12, 10, 5),
    )
    accent = tokens.colors["green"].lstrip("#")
    for name, size, before, after in heading_tokens:
        style = document.styles[name]
        _style_font(style, "Calibri", size, accent)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _compact(header, after=0, line_spacing=1.0)
    _add_static(header, "WEBSITE CONTENT / WORKING FILE", font="Calibri", size=8, bold=True, color=tokens.colors["muted"])
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _compact(footer, after=0, line_spacing=1.0)
    _add_static(
        footer,
        "Edit tagged fields, save in Word, then run the shared working-content sync.",
        font="Calibri",
        size=8,
        color=tokens.colors["muted"],
    )


def _add_document_intro(document: Document, field_count: int, tokens: DesignTokens) -> None:
    title = document.add_paragraph()
    _compact(title, after=3, line_spacing=1.0)
    _add_static(
        title,
        "Website content working file",
        font="Source Serif 4",
        size=26,
        bold=True,
        color=tokens.colors["ink"],
    )
    subtitle = document.add_paragraph()
    _compact(subtitle, after=10, line_spacing=1.15)
    _add_static(
        subtitle,
        f"{field_count} canonical text fields · document schema {DOCUMENT_VERSION}",
        font="DM Mono",
        size=9,
        color=tokens.colors["muted"],
    )
    callout = document.add_table(rows=1, cols=1)
    _set_table_geometry(callout, (TABLE_WIDTH_DXA,))
    cell = callout.cell(0, 0)
    _set_cell_margins(cell)
    _set_cell_fill(cell, tokens.colors["paper_deep"])
    paragraph = cell.paragraphs[0]
    _compact(paragraph, before=2, after=2, line_spacing=1.15)
    _add_static(
        paragraph,
        "HOW TO USE  ",
        font="DM Mono",
        size=8,
        bold=True,
        color=tokens.colors["accent_dark"],
    )
    _add_static(
        paragraph,
        "Edit only the tagged copy in the right column. The left column shows the canonical JSON owner and the exact website typography token. Referenced facts and image descriptions sync to their source JSON rather than replacing $source objects.",
        font="Calibri",
        size=9.5,
        color=tokens.colors["ink"],
    )


def _add_group(
    document: Document, group: str, fields: list[WebsiteContentField], tokens: DesignTokens
) -> None:
    heading = document.add_paragraph(style="Heading 1")
    heading.add_run(group)
    table = document.add_table(rows=1, cols=2)
    _set_table_geometry(table, TABLE_WIDTHS_DXA)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for cell, label in zip(table.rows[0].cells, ("SOURCE & FORMAT CONTEXT", "EDITABLE WEBSITE COPY"), strict=True):
        _set_cell_margins(cell)
        _set_cell_fill(cell, tokens.colors["green"])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        _compact(paragraph, before=1, after=1, line_spacing=1.0)
        _add_static(paragraph, label, font="DM Mono", size=7.5, bold=True, color=tokens.colors["white"])

    for content_field in fields:
        left, right = table.add_row().cells
        for cell in (left, right):
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_fill(left, tokens.colors["paper_deep"])
        _add_field_metadata(left, content_field, tokens)
        _add_field_preview(right, content_field, tokens)


def _add_field_metadata(cell, content_field: WebsiteContentField, tokens: DesignTokens) -> None:
    style = _text_style(content_field.style_name, tokens)
    title = cell.paragraphs[0]
    _compact(title, after=3, line_spacing=1.0)
    _add_static(title, content_field.label, font="Calibri", size=9, bold=True, color=tokens.colors["ink"])

    source = cell.add_paragraph()
    _compact(source, after=3, line_spacing=1.05)
    _add_static(source, content_field.source_label, font="Consolas", size=6.8, color=tokens.colors["muted"])

    metadata = cell.add_paragraph()
    _compact(metadata, after=2, line_spacing=1.05)
    _add_static(metadata, _style_summary(content_field.style_name, style), font="Consolas", size=6.8, color=tokens.colors["ink_soft"])

    if len(content_field.usages) > 1:
        usage = cell.add_paragraph()
        _compact(usage, after=2, line_spacing=1.05)
        _add_static(
            usage,
            f"Used in {len(content_field.usages)} places: " + "; ".join(content_field.usages),
            font="Consolas",
            size=6.4,
            color=tokens.colors["muted"],
        )
    if content_field.template:
        template = cell.add_paragraph()
        _compact(template, after=0, line_spacing=1.05)
        _add_static(
            template,
            f"Rendered template: {content_field.template}",
            font="Consolas",
            size=6.4,
            italic=True,
            color=tokens.colors["accent_dark"],
        )


def _add_field_preview(cell, content_field: WebsiteContentField, tokens: DesignTokens) -> None:
    style = _text_style(content_field.style_name, tokens)
    background = tokens.colors["ink"] if _needs_dark_background(style.color) else tokens.colors["white"]
    _set_cell_fill(cell, background)
    paragraph = cell.paragraphs[0]
    _compact(paragraph, before=3, after=3, line_spacing=_word_line_spacing(style.line_height))
    preview = _word_style(style)
    template = content_field.template or "{value}"
    prefix, separator, suffix = template.partition("{value}")
    if not separator:
        prefix, suffix = "", ""
    if prefix:
        _add_static(paragraph, prefix, **preview)
    _add_control(paragraph, content_field.tag, content_field.value, **preview)
    if suffix:
        _add_static(paragraph, suffix, **preview)


def _text_style(name: str, tokens: DesignTokens) -> TextStyle:
    try:
        return tokens.text_styles[name]
    except KeyError as error:
        raise WebsiteWorkingError(f"Missing design token text.{name}") from error


def _style_summary(name: str, style: TextStyle) -> str:
    return (
        f"text.{name} | Font {style.font_family} | Size {style.font_size} | "
        f"Weight {style.font_weight} | Style {style.font_style} | Line {style.line_height} | "
        f"Tracking {style.letter_spacing} | Transform {style.text_transform}"
    )


def _word_style(style: TextStyle) -> dict[str, Any]:
    return {
        "font": _first_font_family(style.font_family),
        "size": min(_css_size_to_points(style.font_size), PREVIEW_SIZE_CAP_PT),
        "bold": _font_weight(style.font_weight) >= 600,
        "italic": style.font_style.casefold() == "italic",
        "color": style.color,
    }


def _first_font_family(value: str) -> str:
    first = value.split(",", 1)[0].strip().strip('"').strip("'")
    return first or "Calibri"


def _font_weight(value: str) -> int:
    try:
        return int(value)
    except ValueError:
        return 700 if value.casefold() in {"bold", "bolder"} else 400


_CSS_SIZE = re.compile(r"(?P<number>\d*\.?\d+)\s*(?P<unit>rem|px|pt)")


def _css_size_to_points(value: str) -> float:
    match = _CSS_SIZE.search(value)
    if not match:
        return 11.0
    number = float(match.group("number"))
    return number * {"rem": 12.0, "px": 0.75, "pt": 1.0}[match.group("unit")]


def _word_line_spacing(value: str) -> float:
    try:
        return max(1.0, min(float(value), 1.65))
    except ValueError:
        return 1.0


def _needs_dark_background(color: str) -> bool:
    value = color.lstrip("#")
    if len(value) == 3:
        value = "".join(character * 2 for character in value)
    if len(value) < 6:
        return False
    red, green, blue = (int(value[index : index + 2], 16) for index in (0, 2, 4))
    luminance = 0.2126 * red + 0.7152 * green + 0.0722 * blue
    return luminance > 165


def _style_font(style, font: str, size: float, color: str) -> None:
    style.font.name = font
    style._element.rPr.rFonts.set(qn("w:ascii"), font)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color.lstrip("#"))


def _add_control(paragraph, tag: str, value: str, **style: Any) -> None:
    run = paragraph.add_run(value)
    _style_run(run, **style)
    control = OxmlElement("w:sdt")
    properties = OxmlElement("w:sdtPr")
    tag_element = OxmlElement("w:tag")
    tag_element.set(qn("w:val"), tag)
    alias_element = OxmlElement("w:alias")
    alias_element.set(qn("w:val"), tag)
    lock = OxmlElement("w:lock")
    lock.set(qn("w:val"), "sdtLocked")
    properties.extend((tag_element, alias_element, lock))
    content = OxmlElement("w:sdtContent")
    paragraph._p.remove(run._r)
    content.append(run._r)
    control.extend((properties, content))
    paragraph._p.append(control)


def _add_static(paragraph, text: str, **style: Any) -> None:
    run = paragraph.add_run(text)
    _style_run(run, **style)


def _style_run(
    run,
    *,
    font: str = "Calibri",
    size: float = 11,
    bold: bool = False,
    italic: bool = False,
    color: str = "#1F2933",
) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))


def _compact(paragraph, *, before: float = 0, after: float = 0, line_spacing: float = 1.0) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def _set_table_geometry(table, widths: tuple[int, ...]) -> None:
    table.autofit = False
    table.allow_autofit = False
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")
    for grid_column, value in zip(table._tbl.tblGrid.gridCol_lst, widths, strict=True):
        grid_column.set(qn("w:w"), str(value))
    for row in table.rows:
        for cell, value in zip(row.cells, widths, strict=True):
            cell_width = cell._tc.tcPr.tcW
            cell_width.set(qn("w:w"), str(value))
            cell_width.set(qn("w:type"), "dxa")
    _set_table_borders(table)


def _set_table_borders(table) -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "D9E0E6")


def _set_cell_margins(cell) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (
        ("top", CELL_MARGIN_VERTICAL_DXA),
        ("bottom", CELL_MARGIN_VERTICAL_DXA),
        ("start", CELL_MARGIN_HORIZONTAL_DXA),
        ("end", CELL_MARGIN_HORIZONTAL_DXA),
    ):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_cell_fill(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    for existing in properties.findall(qn("w:shd")):
        properties.remove(existing)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color.lstrip("#"))
    properties.append(shading)
