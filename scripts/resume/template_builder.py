"""Create the retained two-page Word template used by the resume pipeline."""

from __future__ import annotations

import argparse
from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# These source-layout colours are converted to the shared design-token palette
# in resume.theme when the public DOCX is built.
NAVY = "18324A"
INK = "1F2933"
MUTED = "5E6C78"
LINE = "D9E0E6"
WHITE = "FFFFFF"
USABLE_WIDTH_DXA = 10800


def create_resume_template(output_path: Path, control_tags: set[str] | None = None) -> None:
    """Create a compact editable template with stable content-control tags."""

    tags = control_tags or _default_template_tags()
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    _configure_styles(document)
    _add_footer(section.footer)
    _add_page_one(document, tags)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    _add_page_two(document, tags)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)

    bullet = document.styles["List Bullet"]
    bullet.font.name = "Arial"
    bullet._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    bullet.font.size = Pt(9.2)


def _add_page_one(document: Document, tags: set[str]) -> None:
    _add_header(document, "CONTACT_NAME", "CONTACT_LOCATION", "CONTACT_PHONE", "CONTACT_EMAIL", "CONTACT_WEBSITE")
    paragraph = document.add_paragraph()
    _compact(paragraph, after=4)
    _add_control(paragraph, "PROFILE_SUMMARY", "Professional summary", size=8.5)

    _section_heading(document, "EDUCATION")
    for tag_prefix, bullets in _education_layout(tags):
        _add_education_line(document, tag_prefix, bullets)

    _section_heading(document, "CORE CAPABILITIES")
    _add_skill_grid(document)

    _section_heading(document, "EXPERIENCE")
    for index, bullets in _entry_layout(tags, "EXP"):
        _add_entry(document, "EXP", index, bullets)

    _section_heading(document, "LEADERSHIP")
    for index, bullets in _entry_layout(tags, "LEAD"):
        _add_entry(document, "LEAD", index, bullets)

    _section_heading(document, "COMMUNITY INVOLVEMENT")
    for index, bullets in _entry_layout(tags, "COMM"):
        _add_entry(document, "COMM", index, bullets)

    _section_heading(document, "RECOGNITION")
    for index, bullets in _entry_layout(tags, "RECOG"):
        _add_entry(document, "RECOG", index, bullets)


def _add_education_line(document: Document, tag_prefix: str, bullets: int) -> None:
    education = document.add_paragraph()
    _compact(education, after=0)
    _add_control(education, f"{tag_prefix}_INSTITUTION", "Institution", bold=True, size=8.6)
    _add_static(education, "  |  ", size=8.6)
    _add_control(education, f"{tag_prefix}_DEGREE", "Program and location", size=8.6)
    _add_static(education, "  |  ", size=8.6)
    _add_control(education, f"{tag_prefix}_DATES", "Dates", bold=True, color=MUTED, size=8.6)
    for bullet_index in range(1, bullets + 1):
        bullet = document.add_paragraph(style="List Bullet")
        _compact(bullet, after=1)
        bullet.paragraph_format.left_indent = Inches(0.17)
        bullet.paragraph_format.first_line_indent = Inches(-0.1)
        _add_control(bullet, f"{tag_prefix}_BULLET{bullet_index}", "Education detail", size=8.4)


def _add_page_two(document: Document, tags: set[str]) -> None:
    header = document.add_paragraph()
    _compact(header, after=7)
    _add_control(header, "PAGE2_NAME", "Name", color=NAVY, bold=True, size=14)
    _add_static(header, "  |  PROJECT PORTFOLIO", bold=True, size=10)
    website = document.add_paragraph()
    _compact(website, after=4)
    website.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_control(website, "PAGE2_WEBSITE", "Portfolio website", color=MUTED, bold=True, size=7.8)

    _section_heading(document, "PROJECT PORTFOLIO")
    for index, bullets in _entry_layout(tags, "PROJECT"):
        _add_entry(document, "PROJECT", index, bullets)

    _section_heading(document, "TECHNICAL SKILLS")
    _add_technical_skills(document)

    callout = document.add_paragraph()
    _compact(callout, before=4, after=0)
    _add_static(callout, "Expanded case studies and technical work: ", italic=True, color=MUTED, size=7.8)
    _add_control(callout, "PORTFOLIO_URL", "Portfolio website", bold=True, color=NAVY, size=7.8)


def _add_header(document: Document, name_tag: str, location_tag: str, phone_tag: str, email_tag: str, website_tag: str) -> None:
    name = document.add_paragraph()
    _compact(name, after=0)
    _add_control(name, name_tag, "Name", color=NAVY, bold=True, size=19)

    contact = document.add_paragraph()
    _compact(contact, after=3)
    _add_control(contact, location_tag, "Location", color=MUTED, bold=True, size=8.1)
    _add_static(contact, "  |  ", color=MUTED, size=8.1)
    _add_control(contact, phone_tag, "Phone", color=MUTED, size=8.1)
    _add_static(contact, "  |  ", color=MUTED, size=8.1)
    _add_control(contact, email_tag, "Email", color=MUTED, size=8.1)
    _add_static(contact, "  |  ", color=MUTED, size=8.1)
    _add_control(contact, website_tag, "Website", color=MUTED, size=8.1)


def _add_entry(document: Document, prefix: str, index: int, bullets: int) -> None:
    if prefix == "PROJECT":
        category = document.add_paragraph()
        _compact(category, before=5, after=1)
        _add_control(
            category,
            f"{prefix}{index}_CATEGORY",
            "Project category",
            bold=True,
            color=NAVY,
            size=8.2,
        )
    title = document.add_paragraph()
    _compact(title, before=2 if prefix == "PROJECT" else 4, after=0)
    _add_control(title, f"{prefix}{index}_TITLE", "Role or title", bold=True, color=NAVY, size=10.2)
    # Date fields are optional for community and recognition records. A visual
    # spacer keeps populated dates readable without leaving a dangling divider
    # when a neutral record has no date yet.
    _add_static(title, "    ", color=MUTED, size=9.0)
    _add_control(title, f"{prefix}{index}_DATES", "Dates", bold=True, color=MUTED, size=9.0)

    metadata = document.add_paragraph()
    _compact(metadata, after=0)
    _add_control(metadata, f"{prefix}{index}_META", "Organization and location", italic=True, color=MUTED, size=8.8)

    for bullet in range(1, bullets + 1):
        paragraph = document.add_paragraph(style="List Bullet")
        _compact(paragraph, after=0)
        paragraph.paragraph_format.left_indent = Inches(0.17)
        paragraph.paragraph_format.first_line_indent = Inches(-0.1)
        _add_control(paragraph, f"{prefix}{index}_BULLET{bullet}", "Key detail", size=9.2)


def _add_skill_grid(document: Document) -> None:
    table = document.add_table(rows=2, cols=3)
    _set_table_geometry(table, (3600, 3600, 3600))
    for index, cell in enumerate((cell for row in table.rows for cell in row.cells), start=1):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _compact(paragraph, before=2, after=2)
        _add_control(paragraph, f"GENERAL_SKILL_{index}", "Capability", bold=True, size=7.6)


def _add_technical_skills(document: Document) -> None:
    rows = (
        ("CAD & DESIGN", "TECH_CAD"),
        ("VEHICLE & ELECTRICAL", "TECH_ELECTRICAL"),
        ("DATA & AUTOMATION", "TECH_DATA"),
        ("ENGINEERING SOFTWARE", "TECH_ENGINEERING_SOFTWARE"),
        ("FABRICATION & QA", "TECH_FABRICATION"),
    )
    table = document.add_table(rows=len(rows), cols=2)
    _set_table_geometry(table, (2450, 8350))
    for row, (label, tag) in zip(table.rows, rows, strict=True):
        label_cell, content_cell = row.cells
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        label_paragraph = label_cell.paragraphs[0]
        _compact(label_paragraph, before=1, after=1)
        _add_static(label_paragraph, label, bold=True, color=NAVY, size=7.2)
        content = content_cell.paragraphs[0]
        _compact(content, before=1, after=1)
        _add_control(content, tag, "Skills", size=7.2)


def _section_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    _compact(paragraph, before=8, after=2)
    _set_paragraph_shading(paragraph, NAVY)
    _set_paragraph_border(paragraph, NAVY)
    _add_static(paragraph, text, bold=True, color=WHITE, size=9.4)


def _add_footer(footer) -> None:
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _compact(paragraph, before=0, after=0)
    _add_static(paragraph, "RESUME / PROJECT PORTFOLIO", bold=True, color=MUTED, size=7.0)


def _add_control(paragraph, tag: str, placeholder: str, **style) -> None:
    run = paragraph.add_run(placeholder)
    _style_run(run, **style)
    sdt = OxmlElement("w:sdt")
    properties = OxmlElement("w:sdtPr")
    tag_element = OxmlElement("w:tag")
    tag_element.set(qn("w:val"), tag)
    alias_element = OxmlElement("w:alias")
    alias_element.set(qn("w:val"), tag)
    properties.append(tag_element)
    properties.append(alias_element)
    content = OxmlElement("w:sdtContent")
    paragraph._p.remove(run._r)
    content.append(run._r)
    sdt.append(properties)
    sdt.append(content)
    paragraph._p.append(sdt)


def _add_static(paragraph, text: str, **style) -> None:
    run = paragraph.add_run(text)
    _style_run(run, **style)


def _style_run(run, *, bold: bool = False, italic: bool = False, color: str = INK, size: float = 8.3) -> None:
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def _compact(paragraph, *, before: float = 0, after: float = 0) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.0


def _set_table_geometry(table, widths: tuple[int, ...]) -> None:
    table.autofit = False
    table.allow_autofit = False
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "0")
    indent.set(qn("w:type"), "dxa")
    for grid_column, value in zip(table._tbl.tblGrid.gridCol_lst, widths, strict=True):
        grid_column.set(qn("w:w"), str(value))
    for row in table.rows:
        for cell, value in zip(row.cells, widths, strict=True):
            cell_width = cell._tc.tcPr.tcW
            cell_width.set(qn("w:w"), str(value))
            cell_width.set(qn("w:type"), "dxa")


def _set_paragraph_shading(paragraph, colour: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), colour)
    paragraph._p.get_or_add_pPr().append(shading)


def _set_paragraph_border(paragraph, colour: str) -> None:
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "1")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), colour)
    borders.append(bottom)
    paragraph._p.get_or_add_pPr().append(borders)


def _education_layout(tags: set[str]) -> list[tuple[str, int]]:
    prefixes = {
        match.group(1)
        for tag in tags
        if (match := re.fullmatch(r"(EDU\d*)_INSTITUTION", tag))
    }
    return [
        (prefix, _bullet_count(tags, prefix))
        for prefix in sorted(prefixes, key=lambda value: 1 if value == "EDU" else int(value.removeprefix("EDU")))
    ]


def _entry_layout(tags: set[str], prefix: str) -> list[tuple[int, int]]:
    indices = sorted(
        int(match.group(1))
        for tag in tags
        if (match := re.fullmatch(rf"{prefix}(\d+)_TITLE", tag))
    )
    return [(index, _bullet_count(tags, f"{prefix}{index}")) for index in indices]


def _bullet_count(tags: set[str], entry_key: str) -> int:
    indices = [
        int(match.group(1))
        for tag in tags
        if (match := re.fullmatch(rf"{entry_key}_BULLET(\d+)", tag))
    ]
    return max(indices, default=0)


def _default_template_tags() -> set[str]:
    tags = {
        "CONTACT_NAME", "CONTACT_LOCATION", "PROFILE_SUMMARY", "CONTACT_PHONE", "CONTACT_EMAIL", "CONTACT_WEBSITE",
        *(f"GENERAL_SKILL_{index}" for index in range(1, 7)),
        "PAGE2_NAME", "PAGE2_WEBSITE", "TECH_CAD", "TECH_ELECTRICAL", "TECH_DATA",
        "TECH_ENGINEERING_SOFTWARE", "TECH_FABRICATION", "PORTFOLIO_URL",
    }
    for index in range(1, 4):
        education = "EDU" if index == 1 else f"EDU{index}"
        tags.update({f"{education}_INSTITUTION", f"{education}_DEGREE", f"{education}_DATES"})
        tags.update(f"{education}_BULLET{bullet}" for bullet in range(1, 3))
    for prefix, count, bullets in (("EXP", 3, 5), ("LEAD", 3, 2), ("COMM", 2, 2), ("RECOG", 4, 2), ("PROJECT", 8, 5)):
        for index in range(1, count + 1):
            entry = f"{prefix}{index}"
            tags.update({f"{entry}_TITLE", f"{entry}_META", f"{entry}_DATES"})
            if prefix == "PROJECT":
                tags.add(f"{entry}_CATEGORY")
            tags.update(f"{entry}_BULLET{bullet}" for bullet in range(1, bullets + 1))
    return tags


def main() -> int:
    parser = argparse.ArgumentParser(description="Rebuild the retained two-page resume Word template")
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    create_resume_template(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
